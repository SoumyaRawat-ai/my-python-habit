from docx import Document

def move_section(input_docx_path, output_docx_path, source_section_num, dest_section_num):
    # Open the input Word document
    doc = Document(input_docx_path)

    # Collect all the paragraphs
    sections = []
    current_section = []
    for paragraph in doc.paragraphs:
        current_section.append(paragraph)
        if paragraph.text.strip() == '':  # Use an empty line as section delimiter
            sections.append(current_section)
            current_section = []
    if current_section:
        sections.append(current_section)

    # Adjust for zero-based indexing
    source_section_num -= 1
    dest_section_num -= 1

    # Get the section to move
    section_to_move = sections.pop(source_section_num)

    # Insert the section at the destination position
    sections.insert(dest_section_num, section_to_move)

    # Create a new Word document with the rearranged sections
    new_doc = Document()
    for section in sections:
        for paragraph in section:
            new_doc.add_paragraph(paragraph.text)

    # Save the output Word document
    new_doc.save(output_docx_path)

input_docx = r"C:\Users\soumy\Downloads\557203348-Our-Project-Report-Online-Food-Ordering.docx"
output_docx = r"C:\Users\soumy\Downloads\557203348-Our-Project-Report-Online-Food-Ordering_Rearranged.docx"
move_section(input_docx, output_docx, 34, 1)  # Example: move 3rd section to 1st position
